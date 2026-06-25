"""Word .docx 解析：段落文本 + 表格（→markdown）+ 内嵌图片。

为简化，按「段落 → 表格 → 图片」的顺序产出块（不严格还原文档内排版顺序）；
对 RAG 检索而言内容是否被收录比块间顺序更重要。
"""
from io import BytesIO

from app.services.document_parsing.base import (
    Block,
    ImageBlock,
    TextBlock,
    UnsupportedFormatError,
    table_to_markdown,
)


def parse_docx(raw: bytes) -> list[Block]:
    try:
        import docx
    except ImportError as exc:
        raise UnsupportedFormatError("服务端未安装 python-docx，无法解析 .docx") from exc

    document = docx.Document(BytesIO(raw))
    blocks: list[Block] = []

    para_text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    if para_text.strip():
        blocks.append(TextBlock(para_text))

    for table in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        md = table_to_markdown(rows)
        if md:
            blocks.append(TextBlock(md))

    for rel in document.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            part = rel.target_part
            ext = part.partname.ext.lstrip(".") or "png"
            blocks.append(ImageBlock(part.blob, ext))
        except Exception:  # noqa: BLE001 个别图片取不到不影响其余内容
            continue

    return blocks
