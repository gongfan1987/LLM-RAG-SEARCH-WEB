"""app/services/document_parsing 单测：分派、各格式解析、表格序列化。

docx/xlsx/pdf 用对应库在内存中生成最小样例，不依赖磁盘文件，也不触达网络。
"""
from io import BytesIO

import pytest

from app.services.document_parsing import (
    ImageBlock,
    TextBlock,
    UnsupportedFormatError,
    parse_document,
)
from app.services.document_parsing.base import table_to_markdown


# ---------- 分派与表格工具 ----------

def test_txt_md_解析为文本块():
    assert parse_document("a.txt", "你好世界".encode("utf-8")) == [TextBlock("你好世界")]
    assert parse_document("a.md", b"# title")[0].text == "# title"


def test_空文本返回空列表():
    assert parse_document("a.txt", b"   \n ") == []


def test_老格式doc给出明确报错():
    with pytest.raises(UnsupportedFormatError) as exc:
        parse_document("legacy.doc", b"x")
    assert ".docx" in str(exc.value)


def test_未知格式报错():
    with pytest.raises(UnsupportedFormatError):
        parse_document("a.zip", b"x")


def test_表格序列化为markdown():
    md = table_to_markdown([["姓名", "年龄"], ["张三", "18"]])
    assert md.splitlines()[0] == "| 姓名 | 年龄 |"
    assert "| --- | --- |" in md
    assert "| 张三 | 18 |" in md


def test_表格行列不齐自动补齐():
    md = table_to_markdown([["a", "b", "c"], ["x"]])
    assert md.splitlines()[-1] == "| x |  |  |"


# ---------- docx ----------

def test_docx_解析段落与表格():
    import docx

    document = docx.Document()
    document.add_paragraph("第一段内容")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "列A"
    table.rows[0].cells[1].text = "列B"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"
    buf = BytesIO()
    document.save(buf)

    blocks = parse_document("d.docx", buf.getvalue())
    texts = [b.text for b in blocks if isinstance(b, TextBlock)]
    assert any("第一段内容" in t for t in texts)
    assert any("| 列A | 列B |" in t for t in texts)


# ---------- xlsx ----------

def test_xlsx_每个工作表转markdown表格():
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "数据"
    ws.append(["产品", "销量"])
    ws.append(["A", 100])
    buf = BytesIO()
    wb.save(buf)

    blocks = parse_document("s.xlsx", buf.getvalue())
    assert len(blocks) == 1 and isinstance(blocks[0], TextBlock)
    assert "工作表「数据」" in blocks[0].text
    assert "| 产品 | 销量 |" in blocks[0].text


# ---------- pdf ----------

def test_pdf_解析出页面文本():
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello PDF content")
    raw = doc.tobytes()
    doc.close()

    blocks = parse_document("p.pdf", raw)
    assert any(isinstance(b, TextBlock) and "Hello PDF content" in b.text for b in blocks)


def test_pdf_注入vl_extract在无表页面不调用vl且走经典提取():
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "plain text without any table here")
    raw = doc.tobytes()
    doc.close()

    calls = {"n": 0}

    def fake_vl(image_bytes):
        calls["n"] += 1
        return "VL_SENTINEL"

    blocks = parse_document("p.pdf", raw, vl_extract=fake_vl)
    assert calls["n"] == 0  # 无表格页面不应触发 VL
    assert any(isinstance(b, TextBlock) and "plain text" in b.text for b in blocks)


def test_导出的块类型仅文本或图片():
    blocks = parse_document("a.txt", b"text")
    assert all(isinstance(b, (TextBlock, ImageBlock)) for b in blocks)
